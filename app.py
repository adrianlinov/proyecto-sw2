from flask import Flask, render_template, request, redirect, session, json, send_file
from flask import Flask
from werkzeug.utils import secure_filename
from database import db, Cliente, Usuario, Valor, Cargo, Categoria, Campo, Servicio, Ocupacion, Rubro, Clientescomite, EstadoSol, VotosComite
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import desc
from sqlalchemy import asc, desc
from docx import Document
from io import BytesIO

app = Flask(__name__)
app.secret_key = "a"
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config[
    'SQLALCHEMY_DATABASE_URI'] = 'postgresql://rnhsdeedutqfbo:9b65509862324147e4b058be064ef659f480885f77c4313fe8a356adc8df3376@ec2-34-195-143-54.compute-1.amazonaws.com:5432/da2uhkjdrgk208'
app.debug = True
db.init_app(app)

with app.app_context():
    db.create_all()

app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False


@app.route('/')
def indexEndPoint():
    categorias = Categoria.query.all()
    return render_template('index.html', categorias=categorias)


@app.route('/formulario', methods=['GET'])
def formularioEndPoint():
    ocupaciones = Ocupacion.query.all()
    rubros = Rubro.query.all()
    all_data = Servicio.query.all()
    campos = Campo.query.all()
    return render_template('/formulario.html',
                           ocupaciones=ocupaciones,
                           rubros=rubros,
                           servicios=all_data,
                           campos=campos)


@app.route('/editar_servicios', methods=['GET'])
def editarServiciosEndPoint():
    if "userId" in session:
        userLogueado = db.session.query(Usuario).filter(
            Usuario.id == session["userId"])[0]
        if userLogueado.cargo_id == 1:

            all_data = Servicio.query.all()
            all_categorias = Categoria.query.all()
            return render_template('editarServicios.html',
                                   servicios=all_data,
                                   categorias=all_categorias, UserLogueado=userLogueado)
        else:
            return redirect("/admin")
    else:
        return redirect("/admin")


@app.route('/editar_servicios', methods=['POST'])
def editarServicioEndPoint():
    userLogueado = db.session.query(Usuario).filter(
        Usuario.id == session["userId"])[0]
    if userLogueado.cargo_id == 1:
        nombre = request.form["Nombre"]
        descripcion = request.form["Descripcion"]
        categoria_id = request.form["Categoria"]
        servicio_id = request.form["Servicio_id"]
        servicio = db.session.query(Servicio).filter(
            Servicio.id == servicio_id)[0]
        servicio.nombre = nombre
        servicio.descripcion = descripcion
        servicio.categoria_id = categoria_id
        db.session.add(servicio)
        db.session.commit()
        return redirect('/editar_servicios')
    else:
        return redirect("/admin")


@app.route('/comercial', methods=['GET'])
def operacionesRegistradasEndPoint():
    if "userId" in session:
        userLogueado = db.session.query(Usuario).filter(
            Usuario.id == session["userId"])[0]
        all_data = Cliente.query.all()
        print(all_data)
        print(userLogueado)
        for row in all_data:
            print(row.servicio)
        all_data_usuario = db.session.query(Usuario).filter(
            Usuario.cargo_id == 4)
        print(all_data_usuario)
        return render_template('comercial.html',
                               Cliente=all_data,
                               Usuario=all_data_usuario,
                               UserLogueado=userLogueado)
    else:
        return redirect("/admin")


@app.route('/editar_usuario')
def crearUsuario():
    if "userId" in session:
        userLogueado = db.session.query(Usuario).filter(
            Usuario.id == session["userId"])[0]
        if userLogueado.cargo_id == 1:
            userLogueado = db.session.query(Usuario).filter(
                Usuario.id == session["userId"])[0]
            all_data = Usuario.query.all()
            all_data_cargo = Cargo.query.all()
            print(all_data_cargo)
            return render_template('crear_usuario.html',
                                   UserLogueado=userLogueado,
                                   Usuarios=all_data,
                                   Cargos=all_data_cargo)
        else:
            return redirect("/admin")
    else:
        return redirect("/admin")


@app.route('/buscarUsuario')
def BuscarUsuario():
    if "userId" in session:
        userLogueado = db.session.query(Usuario).filter(
            Usuario.id == session["userId"])[0]
        usuario = db.session.query(Usuario).filter(
            Usuario.id == request.args.get("user_id"))[0]
        all_data_cargo = Cargo.query.all()
        print(usuario)
        return render_template('editar_usuario.html',
                               usuario=usuario,
                               UserLogueado=userLogueado,
                               Cargos=all_data_cargo)
    else:
        return redirect("/admin")


@app.route('/modificar_usuario', methods=['POST '])
def ModificarUser():
    userLogueado = db.session.query(Usuario).filter(
        Usuario.id == session["userId"])[0]
    if userLogueado.cargo_id == 1:
        id = request.form["id"]
        usuario = request.form["usuario"]
        clave = request.form["clave"]
        cargo = request.form["Id_Cargo"]
        user = db.session.query(Usuario).filter(Usuario.id == id)[0]
        user.usuario = usuario
        user.contrasena = clave
        user.cargo_id = cargo
        db.session.add(user)
        db.session.commit()
        return redirect('/editar_usuario')





@app.route('/EliminarUsuario')
def eliminarUsuario():
    if "userId" in session:
        userLogueado = db.session.query(Usuario).filter(
            Usuario.id == session["userId"])[0]
        if userLogueado.cargo_id == 1:
            userId = request.args.get("user_id")
            print("el user id es", userId)
            Usuario.query.filter_by(id=userId).delete()
            db.session.commit()
            return redirect('/editar_usuario')
        else:
            return redirect("/admin")
    else:
        return redirect("/admin")


@app.route('/editar_usuario', methods=['POST'])
def obtenerUsuario():
    if "userId" in session:
        userLogueado = db.session.query(Usuario).filter(
            Usuario.id == session["userId"])[0]
        if userLogueado.cargo_id == 1:
            all_data = Usuario.query.all()
            usuario = request.form['nombre']
            clave = request.form['clave']
            Id_Cargo = request.form['Id_Cargo']
            print(Id_Cargo)
            print(Id_Cargo)
            nuevo = Usuario(usuario=usuario,
                            contrasena=clave,
                            cargo_id=Id_Cargo)
            db.session.add(nuevo)
            db.session.commit()
            return redirect('/editar_usuario')
        else:
            return redirect("/admin")
    else:
        return redirect("/admin")



@app.route('/guardarmontoaceptado', methods=['POST'])
def guardarmontoaceptado():
    if "userId" in session:
        userLogueado = db.session.query(Usuario).filter(
            Usuario.id == session["userId"])[0]
        if userLogueado.cargo_id == 5:
            print("entro al comite")
            id_cliente = request.form['idCliente']
            result = request.form['resultado']
            usercomite = userLogueado.id 
            flgvotoemitido = "S"
            usuarioanalista = request.form['idUsuarioAnalista']

            if db.session.query(VotosComite).filter(VotosComite.idCliente == id_cliente, VotosComite.idComite == userLogueado.id).count() <= 0:
                nuevo = VotosComite(
                idCliente = id_cliente, 
                idComite = usercomite, 
                resultado = result, 
                flg_votoemitido = flgvotoemitido, 
                idAnalista = usuarioanalista
                )
                db.session.add(nuevo)
                db.session.commit()
                
            else:
                votoPrevio = db.session.query(VotosComite).filter(VotosComite.idCliente == id_cliente, VotosComite.idComite == userLogueado.id)[0]
                votoPrevio.resultado = result
                db.session.add(votoPrevio)
                db.session.commit()
                


            

            cliente = db.session.query(Cliente).filter(Cliente.id == id_cliente)[0]
            cantidadDeVotosAlCliente = db.session.query(VotosComite).filter(VotosComite.idCliente == id_cliente).count()
            cantidadDeUsuariosComite = db.session.query(Usuario).filter(Usuario.cargo_id == 5).count()

            if cantidadDeVotosAlCliente >= cantidadDeUsuariosComite:
                # se completo la cantidad de votos necesarios para validar
                numeroDeVotosAprobados = db.session.query(VotosComite).filter(VotosComite.idCliente == id_cliente, VotosComite.resultado == "APROBADO").count()
                numeroDeVotosRechazado = db.session.query(VotosComite).filter(VotosComite.idCliente == id_cliente, VotosComite.resultado == "RECHAZADO").count()
                if numeroDeVotosAprobados > numeroDeVotosRechazado:
                    
                    cliente.aprobado = "true"
                    db.session.add(cliente)
                    db.session.commit()
                    # enviar al area legal
                    return redirect('/comite')
                else: 
                    pass
                    cliente.aprobado = "false"
                    db.session.add(cliente)
                    db.session.commit()
                    # rechazar usuario
                    return redirect('/comite')

            else:
                return redirect('/comite')
        else:
            return redirect("/admin")
    else:
        return redirect("/admin")

@app.route('/informacion-personal')
def editarInformacionPersonalEndPoint():
    if "userId" in session:
        userLogueado = db.session.query(Usuario).filter(
            Usuario.id == session["userId"])[0]
        cliente = db.session.query(Cliente).filter(
            Cliente.id == request.args.get("cliente_id"))[0]
        valores = db.session.query(Valor).filter(
            Valor.cliente_id == request.args.get("cliente_id"))
        print(valores)
        return render_template('data-personal.html',
                               cliente=cliente,
                               valores=valores, UserLogueado=userLogueado)
    else:
        session.pop("userId", None)
        return redirect("/admin")


@app.route('/nuevo_servicio', methods=['POST'])
def nuevoServicioEndPoint():
    nombre = request.form["Nombre"]
    descripcion = request.form["Descripcion"]
    categoria_id = request.form["Categoria"]
    nuevoServicio = Servicio(nombre=nombre,
                             descripcion=descripcion,
                             categoria_id=categoria_id)
    db.session.add(nuevoServicio)
    db.session.commit()
    print(request.form)

    servicio = db.session.query(Servicio).filter(Servicio.nombre == nombre)[0]
    for key in request.form.keys():
        if key == "Nombre":
            pass
        elif key == "Descripcion":
            pass
        elif key == "Categoria":
            pass
        else:
            print(key)
            nuevoCampo = Campo(nombre=request.form[key],
                               servicio_id=servicio.id)
            db.session.add(nuevoCampo)
            db.session.commit()

    return redirect('/editar_servicios')


@app.route('/nueva_categoria', methods=['POST'])
def nuevaCategoriaEndPoint():
    nombre = request.form["Nombre"]
    categoria = Categoria(nombre=nombre)
    db.session.add(categoria)
    db.session.commit()
    return redirect('/editar_servicios')


@app.route('/formulario', methods=['POST'])
def formularioPost():
    all_data = Servicio.query.all()
    nombre = request.form['Nombre']
    apellido = request.form['Apellido']
    correo = request.form['Correo']
    empresa = request.form['EsEmpresa']
    telefono = request.form['Telefono']
    servicio_id = request.form['servicio_cliente']
    razonSocial = request.form['RazonSocial']
    direccion = request.form['Direccion']
    distrito = request.form['Distrito']
    ciudad = request.form['Ciudad']
    codigoPostal = request.form['CodigoPostal']
    rubroEmpresarial = request.form['RubroEmpresarial']
    montoSolicitado = request.form['MontoSolicitado']
    descripcion = request.form["Mensaje"]

    aux = None
    print(request.form)
    if request.form["EsEmpresa"] == "false":
        documento = request.form['DNI']
        nuevo = Cliente(nombre=nombre,
                        correo=correo,
                        apellido=apellido,
                        esEmpresa=empresa,
                        documento=documento,
                        telefono=telefono,
                        ocupacion_id=request.form['Ocupacion'],
                        servicio_id=servicio_id,
                        direccion=direccion,
                        ciudad=ciudad,
                        distrito=distrito,
                        codigoPostal=codigoPostal,
                        montoSolicitado = montoSolicitado,
                        descripcion = descripcion)
                        
    else:
        documento = request.form['RUC']
        nuevo = Cliente(nombre=nombre,
                        correo=correo,
                        apellido=apellido,
                        esEmpresa=empresa,
                        razonSocial=razonSocial,
                        rubro_id=rubroEmpresarial,
                        documento=documento,
                        telefono=telefono,
                        servicio_id=servicio_id,
                        direccion=direccion,
                        ciudad=ciudad,
                        distrito=distrito,
                        codigoPostal=codigoPostal,
                        montoSolicitado = montoSolicitado)
                        
    print("IDDD")
    print(nuevo.id)
    db.session.add(nuevo)
    db.session.commit()

    clienteGuardado = db.session.query(Cliente).order_by(desc(Cliente.id))[0]
    for key in request.form.keys():
        print(key[0:5])
        print(key[5:])
        if key[0:5] == "campo":
            campoId = key[5:]
            nuevoId = clienteGuardado.id
            nuevoValor = Valor(campo_id=campoId,
                               cliente_id=nuevoId,
                               valor=request.form[key])
            db.session.add(nuevoValor)

    db.session.commit()

    return render_template('formulario.html')


@app.route('/admin', methods=['GET'])
def adminLogInEndPoint():
    return render_template('login.html')


@app.route('/analistas', methods=['GET'])
def analistasEndPoint():
    if "userId" in session:
        UserLogueado = db.session.query(Usuario).filter(Usuario.id == session["userId"])[0]
        clientes_comite = Cliente.query.all()
        lista = []
        for i in clientes_comite:
           lista.append(i.id)
        print(lista)
        all_data = db.session.query(Cliente).filter(
           Cliente.usuario_id == session["userId"], Cliente.id.in_(lista))
        for row in all_data:
           print(row)
           print(row.usuario_id)
        
        user = db.session.query(Usuario).filter(
            Usuario.id == session["userId"])[0]
        return render_template('analistas.html',
                               Cliente=all_data,
                               Usuario=user, UserLogueado=UserLogueado)
        # SI SE LOGUEA EL ADMIN DEBERIA PODER VER A TODOS LOS CLINTES QUE ESTAN DISPONIBLES PARA ANALISIS
    else:
        session.pop("userId", None)
        return redirect("/admin")


@app.route('/comite', methods=['GET'])
def ListarUsuariosEnviadosComite():
    if "userId" in session:
        UserLogueado = db.session.query(Usuario).filter(Usuario.id == session["userId"])[0]
        clientes_comite = db.session.query(Cliente).filter(Cliente.aprobado == None)
        #all_data = db.session.query(Cliente).filter(Cliente.id.in_(lista))
        #all_data = db.session.query(Cliente).filter(Cliente.flg_comite == "S", ~Cliente.id.in_(lista))
        votos = VotosComite.query.all() 
        return render_template('comite.html', Cliente=clientes_comite, VotosEmitidos = votos, UserLogueado=UserLogueado)
        # SI SE LOGUEA EL ADMIN DEBERIA PODER VER A TODOS LOS CLIENTES QUE ESTAN DISPONIBLES PARA ANALISIS
    else:
        session.pop("userId", None)
        return redirect("/admin")


@app.route('/GuardarVoto', methods=['POST'])
def GuardarVoto():
    if "userId" in session:
        userLog = db.session.query(Usuario).filter(Usuario.id == session["userId"])[0]
        if (userLog.cargo_id == 5 or userLog.cargo_id == 1):
            userLogueado = session["userId"]
            id_cliente = request.form["idCliente"]
            nuevo = Clientescomite(idComite=userLogueado, idCliente=id_cliente)
            db.session.add(nuevo)
            db.session.commit()
            return redirect('/analistas')
        else:
            session.pop("userId", None)
            return redirect("/admin")
        # SI SE LOGUEA EL ADMIN DEBERIA PODER VER A TODOS LOS CLINTES QUE ESTAN DISPONIBLES PARA ANALISIS
    else:
        session.pop("userId", None)
        return redirect("/admin")


@app.route('/analisis', methods=['GET'])
def analisisEndPoint():
    if "userId" in session:
        userLogueado = db.session.query(Usuario).filter(
            Usuario.id == session["userId"])[0]
        print(userLogueado.cargo_id)
        if (userLogueado.cargo_id == 4 or userLogueado.cargo_id
                == 1) and request.args.get("edit") == "true":
            cliente = db.session.query(Cliente).filter(
                Cliente.id == request.args.get("cliente_id"))[0]
            return render_template('analisis.html',
                                   cliente=cliente,
                                   editar="true", UserLogueado=userLogueado)
        else:
            cliente = db.session.query(Cliente).filter(
                Cliente.id == request.args.get("cliente_id"))[0]
            return render_template('analisis.html',
                                   cliente=cliente,
                                   editar="false", UserLogueado=userLogueado)
    else:
        session.pop("userId", None)
        return redirect("/admin")


@app.route('/analisis', methods=['POST'])
def analisisEndPointPost():
    if "userId" in session:
        userLogueado = db.session.query(Usuario).filter(
            Usuario.id == session["userId"])[0]
        print(userLogueado.cargo_id)
        if (userLogueado.cargo_id == 4 or userLogueado.cargo_id == 1):
            cliente = db.session.query(Cliente).filter(
                Cliente.id == request.form["clienteid"])[0]
            cliente.analisis = request.form["mytextarea"]
            db.session.add(cliente)
            db.session.commit()
            return render_template('analisis.html',
                                   cliente=cliente,
                                   editar="true", UserLogueado=userLogueado)
        else:
            cliente = db.session.query(Cliente).filter(
                Cliente.id == request.form["clienteid"])[0]
            return render_template('analisis.html',
                                   cliente=cliente,
                                   editar="false", UserLogueado=userLogueado)
    else:
        session.pop("userId", None)
        return redirect("/admin")


@app.route('/nosotros', methods=['GET'])
def nosotrosEndPoint():
    return render_template('nosotros.html')


@app.route('/admin', methods=['POST'])
def submit():
    if request.method == 'POST':
        usuario = request.form['user']
        contrasena = request.form['pwduser']

        user = db.session.query(Usuario).filter(
            Usuario.usuario == usuario, Usuario.contrasena == contrasena)
        if (user.count() > 0):
            session["userId"] = user[0].id
            session["userCargo_id"] = user[0].cargo_id
            print("USUARIO LOGEADO: ", session["userId"])
            print("CARGO LOGEADO: ", session["userCargo_id"])
            if (user[0].cargo_id == 1):
                return redirect('/comercial')
            elif (user[0].cargo_id == 2):
                return redirect('/ventas')
            elif (user[0].cargo_id == 3):
                return redirect('/comercial')
            elif (user[0].cargo_id == 4):
                return redirect('/analistas')
            elif (user[0].cargo_id == 5):
                return redirect('/comite')
            elif (user[0].cargo_id == 6):
                return redirect('/legal')
            else:
                session.pop("userId", None)
                return redirect('/admin')

        else:
            session.pop("userId", None)
            return redirect('/admin')


@app.route('/Asignar_Analista', methods=['POST'])
def Asignar_Analista():
    if "userId" in session:
        userLogueado = db.session.query(Usuario).filter(
            Usuario.id == session["userId"])[0]
        if userLogueado.cargo_id == 3 or userLogueado.cargo_id == 1:
            user = db.session.query(Usuario).filter(
                Usuario.id == request.form["Id_Usuario"])[0]
            cliente = db.session.query(Cliente).filter(
                Cliente.id == request.form["Id_Cliente"])[0]
            cliente.usuario = user
            db.session.add(cliente)
            db.session.commit()
            return redirect('/comercial')
        else:
            session.pop("userId", None)
            return redirect('/admin')


@app.route('/EnviarComite')
def Enviar_comite():
    if "userId" in session:
        userLogueado = session["userId"]
        id_cliente = request.args.get("cliente_id")
        print(userLogueado)
        cliente_encontrado = db.session.query(Cliente).filter(Cliente.id == id_cliente)[0]
        cliente_encontrado.flg_comite = "S"
        db.session.add(cliente_encontrado)
        db.session.commit()
        return redirect('/analistas')
    else:
        session.pop("userId", None)
        return redirect('/admin')


@app.route('/legal', methods=['GET'])
def legalEndPoint():
    if "userId" in session:
        userLogueado2 = db.session.query(Usuario).filter(
            Usuario.id == session["userId"])[0]
        userLogueado = session["userId"]
        all_data = db.session.query(Cliente).filter(Cliente.aprobado == "true", Cliente.completado == None)
        print(all_data)
        for row in all_data:
            print(row.completado)
        all_data_usuario = db.session.query(Usuario).filter(
            Usuario.cargo_id == 4)
        return render_template('legal.html', UserLogueado=userLogueado2, Cliente=all_data,
                               Usuario=all_data_usuario)
    else:
        session.pop("userId", None)
        return redirect('/admin')

@app.route('/generar-contrato', methods=['POST'])
def generarContratoPoint():
    print("generar contrato")
    if "userId" in session:
        userLogueado = session["userId"]
        cliente = db.session.query(Cliente).filter(
                Cliente.id == request.form["clienteid"])[0]
        # Debe contar con un IF para seleccionar el tipo de contrato deseado, en caso sea Persona Natural o Juridica
        # Debe hacer el cambio con los parametros de cada cliente
        print(cliente.__dict__.keys())

        document = None
        if cliente.esEmpresa == True:
            document = Document('Contrato - Empresa.docx')  
        else:
            document = Document('Contrato - Persona.docx')
        for key in cliente.__dict__.keys():
            print("{" + key + "}")
            palabra = "{" + key + "}"
            remplazo = str(cliente.__dict__[key])
            for paragraph in document.paragraphs:
                if palabra in paragraph.text:
                    paragraph.text = paragraph.text.replace(palabra, remplazo)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if palabra in paragraph.text:
                                print(paragraph.text)
                                paragraph.text = paragraph.text.replace(palabra, remplazo)

        # El nombre del documento retornado debe ser con el nombre del cliente

        if cliente.esEmpresa == True:
            document.save('Contrato - Empresa - {}.docx'.format(cliente.razonSocial))    
            return send_file('Contrato - Empresa - {}.docx'.format(cliente.razonSocial))
        else:
            document.save('Contrato - Persona - {}.docx'.format(cliente.nombre))    
            return send_file('Contrato - Persona - {}.docx'.format(cliente.nombre))
    else:
        session.pop("userId", None)
        return redirect('/admin')





@app.route('/cargar-contrato', methods=['POST'])
def cargarContrato():
    if "userId" in session:
        userLogueado = session["userId"]
        cliente = db.session.query(Cliente).filter(
                Cliente.id == request.form["clienteid"])[0]
        contratoCargado = request.files["contrato"].read()
        cliente.contrato = contratoCargado
        db.session.add(cliente)
        db.session.commit()
        return redirect('/legal')
    else:
        session.pop("userId", None)
        return redirect('/admin')


@app.route('/descargarContrato', methods=['GET'])
def descargarContrato():
    if "userId" in session:
        userLogueado = session["userId"]
        cliente = db.session.query(Cliente).filter(Cliente.id == request.args.get("clienteid"))[0]

        return send_file(BytesIO(cliente.contrato) ,as_attachment=True, attachment_filename="Contrato - {}.docx".format(cliente.nombre))
    
    else:
        session.pop("userId", None)
        return redirect('/admin')


@app.route('/completarCliente', methods=['GET'])
def completarCliente():
    if "userId" in session: 
        userLogueado = session["userId"]
        cliente = db.session.query(Cliente).filter(Cliente.id == request.args.get("clienteid"))[0]
        cliente.completado = "true"
        db.session.add(cliente)
        db.session.commit()
        return redirect("/legal")
    
    else:
        session.pop("userId", None)
        return redirect('/admin')




@app.route('/completado', methods=['GET'])
def completadoEndPoint():
    if "userId" in session:
        userLogueado2 = db.session.query(Usuario).filter(
            Usuario.id == session["userId"])[0]
        userLogueado = session["userId"]
        all_data = db.session.query(Cliente).filter(Cliente.completado == "true")
        for row in all_data:
            print(row.servicio)
        all_data_usuario = db.session.query(Usuario).filter(
            Usuario.cargo_id == 4)
        return render_template('completado.html', UserLogueado=userLogueado2, Cliente=all_data,
                               Usuario=all_data_usuario)
    else:
        session.pop("userId", None)
        return redirect('/admin')



@app.route('/rechazado', methods=['GET'])
def rechazadoEndPoint():
    if "userId" in session:
        userLogueado2 = db.session.query(Usuario).filter(
            Usuario.id == session["userId"])[0]
        userLogueado = session["userId"]
        all_data = db.session.query(Cliente).filter(Cliente.aprobado == "false")
        for row in all_data:
            print(row.servicio)
        all_data_usuario = db.session.query(Usuario).filter(
            Usuario.cargo_id == 4)
        return render_template('rechazado.html', UserLogueado=userLogueado2, Cliente=all_data,
                               Usuario=all_data_usuario)
    else:
        session.pop("userId", None)
        return redirect('/admin')

@app.route('/inServicio', methods=['GET'])
def inServicioEndpoint():
    servicio = db.session.query(Servicio).filter(Servicio.id == request.args.get("servicioid"))[0]
    print(servicio)
    return render_template("inServicio.html", servicio=servicio)









@app.route("/logout", methods=['GET'])
def logOut():
    session.pop("userId", None)
    return redirect("/admin")


if __name__ == '__main__':
    app.debug = True
    app.run(host='0.0.0.0', port=80)
